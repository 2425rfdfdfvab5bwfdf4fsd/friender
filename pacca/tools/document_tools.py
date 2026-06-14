"""Document tools — create_docx, read_docx, create_xlsx, read_xlsx."""
from __future__ import annotations
from pathlib import Path
from typing import Any


def create_docx(path: str, title: str = "", content: str = "",
                dry_run: bool = False) -> dict:
    p = Path(path)
    if p.exists():
        return {"error": f"File already exists: {path}", "blocked": True}
    if dry_run:
        return {"dry_run": True, "would_create": str(p.resolve())}
    try:
        from docx import Document
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        if content:
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return {"created": str(p.resolve()), "title": title}
    except ImportError:
        return {"error": "python-docx not installed"}
    except Exception as e:
        return {"error": str(e)}


def read_docx(path: str) -> dict:
    p = Path(path)
    if not p.exists():
        return {"error": f"File not found: {path}"}
    try:
        from docx import Document
        doc = Document(str(p))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        content = "\n\n".join(paragraphs)
        return {
            "path": str(p.resolve()),
            "content": content,
            "paragraph_count": len(paragraphs),
            "character_count": len(content),
        }
    except ImportError:
        return {"error": "python-docx not installed"}
    except Exception as e:
        return {"error": str(e)}


def create_xlsx(path: str, sheet_name: str = "Sheet1",
                headers: list[str] | None = None,
                rows: list[list] | None = None,
                dry_run: bool = False) -> dict:
    p = Path(path)
    if p.exists():
        return {"error": f"File already exists: {path}", "blocked": True}
    if dry_run:
        return {"dry_run": True, "would_create": str(p.resolve())}
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name
        if headers:
            ws.append(headers)
        if rows:
            for row in rows:
                ws.append(row)
        p.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(p))
        return {
            "created": str(p.resolve()),
            "sheet": sheet_name,
            "rows": len(rows) if rows else 0,
            "columns": len(headers) if headers else 0,
        }
    except ImportError:
        return {"error": "openpyxl not installed"}
    except Exception as e:
        return {"error": str(e)}


def read_xlsx(path: str, sheet_name: str | None = None,
              max_rows: int = 1000) -> dict:
    p = Path(path)
    if not p.exists():
        return {"error": f"File not found: {path}"}
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                return {"error": f"Sheet '{sheet_name}' not found"}
            ws = wb[sheet_name]
        else:
            ws = wb.active
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            rows.append(list(row))
        return {
            "path": str(p.resolve()),
            "sheet": ws.title,
            "sheets": wb.sheetnames,
            "rows": rows,
            "row_count": len(rows),
            "column_count": ws.max_column,
            "truncated": ws.max_row > max_rows if ws.max_row else False,
        }
    except ImportError:
        return {"error": "openpyxl not installed"}
    except Exception as e:
        return {"error": str(e)}
